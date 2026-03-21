"""
DOCX document extraction and loading.

This module provides the DocxExtractor class for loading and accessing
components of a DOCX document using the python-docx library.
"""

import logging
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph

from .exceptions import InvalidDocxError

logger = logging.getLogger(__name__)


class DocxExtractor:
    """Extracts and provides access to DOCX document components.
    
    This class is responsible for loading a DOCX file and providing
    structured access to its contents (paragraphs, tables, images, etc.).
    
    Attributes:
        file_path: Path to the DOCX file
        document: Loaded Document object from python-docx
        
    Example:
        >>> extractor = DocxExtractor(Path("document.docx"))
        >>> extractor.load_document()
        >>> paragraphs = extractor.get_paragraphs()
    """
    
    def __init__(self, file_path: Path):
        """Initialize the extractor with a file path.
        
        Args:
            file_path: Path to the DOCX file to extract.
        """
        self.file_path = file_path
        self.document: Document = None
        logger.debug(f"DocxExtractor initialized for: {file_path}")
    
    def load_document(self) -> Document:
        """Load the DOCX document.
        
        Returns:
            The loaded Document object.
            
        Raises:
            FileNotFoundError: If the input file doesn't exist.
            InvalidDocxError: If the file is not a valid DOCX document.
        """
        if not self.file_path.exists():
            logger.error(f"File not found: {self.file_path}")
            raise FileNotFoundError(f"Input file not found: {self.file_path}")
        
        try:
            logger.info(f"Loading DOCX document: {self.file_path}")
            self.document = Document(self.file_path)
            logger.info(f"Successfully loaded document with {len(self.document.paragraphs)} paragraphs")
            return self.document
        
        except PackageNotFoundError as e:
            logger.error(f"Invalid DOCX file: {e}")
            raise InvalidDocxError(
                str(self.file_path),
                "File is not a valid DOCX document (not a valid ZIP package)"
            )
        
        except Exception as e:
            logger.error(f"Error loading DOCX document: {e}")
            raise InvalidDocxError(
                str(self.file_path),
                f"Failed to load DOCX document: {str(e)}"
            )
    
    def get_paragraphs(self) -> List[Paragraph]:
        """Get all paragraphs from the document.
        
        Returns:
            List of Paragraph objects from the document.
            
        Raises:
            RuntimeError: If document hasn't been loaded yet.
        """
        if self.document is None:
            raise RuntimeError("Document not loaded. Call load_document() first.")
        
        return self.document.paragraphs
    
    def get_tables(self) -> List[Table]:
        """Get all tables from the document.
        
        Returns:
            List of Table objects from the document.
            
        Raises:
            RuntimeError: If document hasn't been loaded yet.
        """
        if self.document is None:
            raise RuntimeError("Document not loaded. Call load_document() first.")
        
        return self.document.tables
    
    def get_image_parts(self) -> Dict[str, bytes]:
        """Extract image parts from the document.
        
        Returns:
            Dictionary mapping relationship IDs to image binary data.
            
        Raises:
            RuntimeError: If document hasn't been loaded yet.
        """
        if self.document is None:
            raise RuntimeError("Document not loaded. Call load_document() first.")
        
        image_parts = {}
        
        try:
            # Access document part relationships
            for rel in self.document.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_parts[rel.rId] = rel.target_part.blob
                        logger.debug(f"Found image: {rel.rId} -> {rel.target_ref}")
                    except Exception as e:
                        logger.warning(f"Could not extract image {rel.rId}: {e}")
        
        except Exception as e:
            logger.warning(f"Error extracting image parts: {e}")
        
        logger.info(f"Extracted {len(image_parts)} image parts")
        return image_parts
